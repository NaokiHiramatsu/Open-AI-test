import os
import uuid
from flask import Flask, request, jsonify, render_template, session, send_file, url_for
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
import openai
import pandas as pd
from io import BytesIO
from flask_session import Session
import requests
from fpdf import FPDF
from docx import Document
import time
from PyPDF2 import PdfReader

# Flaskアプリの設定
app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config['SESSION_TYPE'] = 'filesystem'
Session(app)

# 環境変数の設定
openai.api_type = "azure"
openai.api_base = os.getenv("OPENAI_API_BASE", "https://example.openai.azure.com")
openai.api_version = "2024-08-01-preview"
openai.api_key = os.getenv("OPENAI_API_KEY", "your-api-key")
response_model = os.getenv("OPENAI_RESPONSE_MODEL", "response-model")

search_service_endpoint = os.getenv("AZURE_SEARCH_ENDPOINT", "https://search-service.azure.com")
search_service_key = os.getenv("AZURE_SEARCH_KEY", "your-search-key")
index_name = os.getenv("AZURE_SEARCH_INDEX_NAME", "your-index-name")

vision_subscription_key = os.getenv("VISION_API_KEY", "your-vision-key")
vision_endpoint = os.getenv("VISION_ENDPOINT", "https://vision.azure.com")

# Azure Search クライアントの設定
try:
    if not all([search_service_endpoint, search_service_key, index_name]):
        raise ValueError("Azure Search の環境変数が正しく設定されていません。")

    search_client = SearchClient(
        endpoint=search_service_endpoint,
        index_name=index_name,
        credential=AzureKeyCredential(search_service_key)
    )
    print("Search client initialized successfully.")
except Exception as e:
    search_client = None
    print(f"SearchClient initialization failed: {e}")

# Azure Search接続確認関数
def check_search_connection():
    try:
        test_response = requests.get(
            f"{search_service_endpoint}/indexes?api-version=2021-04-30-Preview",
            headers={"api-key": search_service_key}
        )
        test_response.raise_for_status()
        print("Azure Search connection is successful.")
    except Exception as e:
        print(f"Failed to connect to Azure Search: {e}")
        return False
    return True

# 一時ファイル保存ディレクトリを作成
SAVE_DIR = "generated_files"
if not os.path.exists(SAVE_DIR):
    os.makedirs(SAVE_DIR)

@app.route('/')
def index():
    session.clear()
    return render_template('index.html', chat_history=[])

@app.route('/process_files_and_prompt', methods=['POST'])
def process_files_and_prompt():
    files = request.files.getlist('files')
    prompt = request.form.get('prompt', '')

    if 'chat_history' not in session:
        session['chat_history'] = []

    try:
        file_data_text = []
        excel_dataframes = []
        extracted_data = []

        for file in files:
            if file and file.filename.endswith('.xlsx'):
                df = pd.read_excel(file, engine='openpyxl')
                excel_dataframes.append(df)
                extracted_data.append({"type": "excel", "data": df.to_dict(orient='records')})
            elif file and file.filename.endswith('.pdf'):
                pdf_text = extract_text_from_pdf(file)
                extracted_data.append({"type": "pdf", "data": pdf_text})
            elif file and file.filename.endswith('.docx'):
                word_text = extract_text_from_word(file)
                extracted_data.append({"type": "word", "data": word_text})
            elif file and file.filename.endswith(('.png', '.jpg', '.jpeg')):
                image_data = file.read()
                image_filename = save_image_to_temp(image_data)
                ocr_text = ocr_image(image_filename)
                extracted_data.append({"type": "image", "data": ocr_text})

        if search_client and check_search_connection():
            try:
                search_results = search_client.search(
                    search_text=prompt,
                    query_type="semantic",
                    semantic_configuration_name="vector-1730110777868-semantic-configuration",
                    select=["chunk", "title", "chunk_id"]
                )
                for result in search_results:
                    extracted_data.append({
                        "type": "search",
                        "data": result
                    })
            except Exception as e:
                extracted_data.append({"type": "error", "data": str(e)})

        response_content, output_format = generate_ai_response_and_format(extracted_data, response_model)

        session['chat_history'].append({
            'user': extracted_data,
            'assistant': response_content
        })

        file_output = None
        if output_format == "xlsx" and excel_dataframes:
            combined_excel = BytesIO()
            with pd.ExcelWriter(combined_excel, engine='openpyxl') as writer:
                for i, df in enumerate(excel_dataframes):
                    df.to_excel(writer, index=False, sheet_name=f"Sheet{i+1}")
            combined_excel.seek(0)
            file_output = combined_excel
        elif output_format == "pdf":
            file_output = generate_pdf(response_content)
        elif output_format == "docx":
            file_output = generate_word(response_content)
        else:
            file_output = BytesIO(response_content.encode('utf-8'))

        temp_filename = f"{uuid.uuid4()}.{output_format}"
        file_path = os.path.join(SAVE_DIR, temp_filename)
        with open(file_path, 'wb') as f:
            f.write(file_output.getvalue())

        download_url = url_for('download_file', filename=temp_filename, _external=True)
        session['chat_history'][-1]['assistant'] += f" <a href='{download_url}' target='_blank'>生成されたファイルをダウンロード</a>"

        return render_template('index.html', chat_history=session['chat_history'])

    except Exception as e:
        print(f"Error during file processing: {e}")
        return jsonify({"error": f"エラーが発生しました: {e}"}), 500

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.abspath(os.path.join(SAVE_DIR, filename))
    if not file_path.startswith(os.path.abspath(SAVE_DIR)):
        return "Invalid file path", 400

    if os.path.exists(file_path):
        mimetype_map = {
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain'
        }
        ext = filename.split('.')[-1]
        mimetype = mimetype_map.get(ext, 'application/octet-stream')
        return send_file(file_path, mimetype=mimetype, as_attachment=True, download_name=filename)
    return "File not found or file is empty", 404

def save_image_to_temp(image_data):
    temp_filename = f"{uuid.uuid4()}.png"
    temp_path = os.path.join(SAVE_DIR, temp_filename)
    with open(temp_path, 'wb') as f:
        f.write(image_data)
    return temp_path

def ocr_image(image_path):
    ocr_url = f"{vision_endpoint}/vision/v3.2/ocr"
    headers = {"Ocp-Apim-Subscription-Key": vision_subscription_key}
    with open(image_path, "rb") as img:
        response = requests.post(ocr_url, headers=headers, files={"file": img})
    response.raise_for_status()
    ocr_results = response.json()
    return "\n".join([" ".join(word['text'] for word in line['words']) for region in ocr_results.get('regions', []) for line in region.get('lines', [])])

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join(page.extract_text() for page in reader.pages)

def extract_text_from_word(file):
    doc = Document(file)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

def generate_pdf(content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, content)
    output = BytesIO()
    pdf.output(output)
    output.seek(0)
    return output

def generate_word(content):
    doc = Document()
    doc.add_paragraph(content)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_ai_response_and_format(input_data, deployment_name):
    messages = [
        {"role": "system", "content": (
            "あなたは、システム内で直接ファイルを生成し、適切な形式（text, Excel, PDF, Word）を判断し生成します。"
            "生成するExcelファイルには、1行目に列名、2行目以降にデータ行を含める必要があります。"
            "ファイルで出力すべき内容以外の文章はテキスト形式で返してください。"
            "必ずファイル形式と内容を判断し、必要に応じて表形式を正しく出力してください。"
            "Flaskの/downloadエンドポイントを使用してリンクをHTML <a>タグで提供してください。"
        )},
        {"role": "user", "content": input_data}
    ]
    retry_count = 0
    while retry_count < 3:  # 最大3回リトライ
        try:
            response = openai.ChatCompletion.create(
                engine=deployment_name, messages=messages, max_tokens=2000
            )
            response_text = response['choices'][0]['message']['content']
            output_format = determine_output_format_from_response(response_text)
            return response_text, output_format
        except openai.error.RateLimitError:
            retry_count += 1
            wait_time = 10 + (retry_count * 5)  # 待機時間を段階的に増やす
            print(f"Rate limit exceeded. Retrying in {wait_time} seconds...")
            time.sleep(wait_time)
    raise Exception("Rate limit exceeded. Please try again later.")

def determine_output_format_from_response(response_content):
    if "excel" in response_content.lower(): return "xlsx"
    if "pdf" in response_content.lower(): return "pdf"
    if "word" in response_content.lower(): return "docx"
    return "txt"

if __name__ == '__main__':
    app.run(debug=True)
